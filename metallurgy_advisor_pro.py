import streamlit as st
import pandas as pd
import numpy as np
from sklearn.ensemble import GradientBoostingRegressor
import io
import joblib
from itertools import product
import matplotlib.pyplot as plt
import seaborn as sns
try:
    from docx import Document
except ImportError:
    st.error("Установите библиотеку python-docx: pip install python-docx")

# --- 1. НАСТРОЙКА ИНТЕРФЕЙСА ---
st.set_page_config(page_title="Ультра Советник Металлурга", layout="wide", page_icon="🏭")
st.title("🏭 Ультра-Советник: От выплавки до проката")

# --- 2. КОНСТАНТЫ И ФУНКЦИИ ---
def engineer_features(df):
    df_new = df.copy()
    if 'Mn' in df_new.columns and 'S' in df_new.columns:
        df_new['Mn/S'] = df_new['Mn'] / df_new['S'].replace(0, 0.0001)
    return df_new

def calculate_feature_direction(models, features, df_ref):
    """Рассчитывает направление влияния каждого параметра на целевые показатели"""
    directions = {}
    
    for target, model in models.items():
        target_directions = {}
        
        # Берём средние значения из данных
        X_mean = df_ref[features].mean().to_dict()
        X_test = pd.DataFrame([X_mean])
        baseline_pred = model.predict(X_test)[0]
        
        for feature in features:
            # Увеличиваем параметр на 10% и смотрим изменение
            X_increased = X_test.copy()
            increase_val = X_mean[feature] * 0.1 if X_mean[feature] != 0 else 0.1
            X_increased[feature] = X_mean[feature] + increase_val
            new_pred = model.predict(X_increased)[0]
            
            # Определяем направление: + (положительное) или - (отрицательное)
            delta = new_pred - baseline_pred
            target_directions[feature] = "↑" if delta > 0 else ("↓" if delta < 0 else "~")
        
        directions[target] = target_directions
    
    return directions

base_features = ['Перегрев МНЛЗ(конечный)', 'Перегрев МНЛЗ (начальный)', 'H', 'О', 'Cr', 'Mn', 'N', 'Al', 'As', 'Mn/S', 'Ti']
termo_features = ['Продолжительность нагрева, мин', 'Факт время закалки ВЗМ, с', 'темп-ра отпуска, °С']
extended_features = base_features + termo_features
targets = ['KV Среднее', 'KV мин', 'К1С Среднее', 'К1С мин']

# Инициализация хранилища моделей
if 'base_models' not in st.session_state: st.session_state.base_models = None
if 'extended_models' not in st.session_state: st.session_state.extended_models = None

# --- 3. БЛОК УПРАВЛЕНИЯ МОДЕЛЯМИ (ОБУЧЕНИЕ И ЗАГРУЗКА) ---
with st.expander("⚙️ УПРАВЛЕНИЕ МОДЕЛЯМИ (Обучить новые или загрузить готовые)", expanded=(st.session_state.base_models is None)):
    st.markdown("Здесь вы можете загрузить исторические данные для обучения нейросетей ИЛИ загрузить уже готовые файлы `.joblib`.")
    tab_train, tab_load = st.tabs(["🏗️ Обучение по историческим данным", "📂 Загрузка готовых файлов моделей"])
    
    with tab_train:
        uploaded_train = st.file_uploader("Загрузите исторический Excel-файл (должен содержать химию, термообработку и 4 мех. свойства)", type=['xlsx', 'xls'])
        if uploaded_train:
            df_train = engineer_features(pd.read_excel(uploaded_train))
            missing = [c for c in extended_features + targets if c not in df_train.columns and c != 'Mn/S']
            if missing:
                st.error(f"В файле не хватает колонок: {', '.join(missing)}")
            else:
                if st.button("Запустить обучение", type="primary"):
                    with st.spinner("Обучаю 8 нейросетей..."):
                        b_models, e_models = {}, {}
                        for t in targets:
                            df_clean_base = df_train[base_features + [t]].dropna()
                            b_models[t] = GradientBoostingRegressor(random_state=42, n_estimators=150, max_depth=4).fit(df_clean_base[base_features], df_clean_base[t])
                            
                            df_clean_ext = df_train[extended_features + [t]].dropna()
                            e_models[t] = GradientBoostingRegressor(random_state=42, n_estimators=150, max_depth=4).fit(df_clean_ext[extended_features], df_clean_ext[t])
                        
                        st.session_state.base_models = b_models
                        st.session_state.extended_models = e_models
                    st.success("✅ Модели обучены! Теперь вы можете скачать их или сразу перейти к анализу.")
                    col_dl1, col_dl2 = st.columns(2)
                    buf_b, buf_e = io.BytesIO(), io.BytesIO()
                    joblib.dump(b_models, buf_b); joblib.dump(e_models, buf_e)
                    col_dl1.download_button("💾 Скачать Базовую", buf_b.getvalue(), "model_base.joblib")
                    col_dl2.download_button("💾 Скачать Расширенную", buf_e.getvalue(), "model_extended.joblib")
    
    with tab_load:
        col_L1, col_L2 = st.columns(2)
        up_b = col_L1.file_uploader("Загрузить model_base.joblib", type=['joblib'])
        up_e = col_L2.file_uploader("Загрузить model_extended.joblib", type=['joblib'])
        if up_b: st.session_state.base_models = joblib.load(up_b)
        if up_e: st.session_state.extended_models = joblib.load(up_e)
        if st.session_state.base_models and st.session_state.extended_models:
            st.success("Модели успешно загружены в память!")

st.sidebar.header("Режимы работы")
mode = st.sidebar.radio("Выберите функцию:", [
    "1. Сталеплавильная (Прогноз + Оптимизация)",
    "2. Советник по прокату (Рекомендация режимов)",
    "3. Комбинированная (Полный анализ)",
    "4. Ручной Симулятор (Песочница)"
])

if not (st.session_state.base_models and st.session_state.extended_models):
    st.warning("⚠️ Пожалуйста, обучите или загрузите модели в блоке сверху ⚙️")
    st.stop()

st.divider()

# =========================================================================
# ФУНКЦИЯ 1: СТАЛЕПЛАВИЛЬНАЯ
# =========================================================================
if mode == "1. Сталеплавильная (Прогноз + Оптимизация)":
    st.header("Функция 1: Прогноз по химии и оптимизация параметра")
    
    # Показываем важность параметров (топ 10)
    with st.expander("📊 Анализ влияния параметров на механические свойства", expanded=False):
        st.subheader("Важность параметров для стали (базовые свойства)")
        st.info("📈 **↑** = повышение параметра улучшает показатель | **↓** = повышение параметра ухудшает показатель | **~** = слабое влияние")
        
        # Собираем важности параметров из всех моделей
        importance_dict = {}
        for target in targets:
            model = st.session_state.base_models[target]
            importance_dict[target] = model.feature_importances_
        
        # Создаём DataFrame с важностями
        importance_df = pd.DataFrame(importance_dict, index=base_features)
        
        # Берём топ 10 параметров по средней важности
        importance_df['Mean'] = importance_df.mean(axis=1)
        top_10_params = importance_df.nlargest(10, 'Mean').drop('Mean', axis=1)
        
        # Рассчитываем направление влияния
        direction_text = pd.DataFrame(index=top_10_params.index, columns=top_10_params.columns, dtype=object)
        try:
            # Создаём df для расчёта средних значений
            df_dummy = pd.DataFrame({col: [0] for col in base_features})
            for col in base_features:
                df_dummy[col] = np.random.uniform(0, 1, 100)
            
            directions = calculate_feature_direction(st.session_state.base_models, base_features, df_dummy)
            
            # Добавляем направления в DataFrame для аннотаций
            for feat in top_10_params.index:
                for target in top_10_params.columns:
                    val = top_10_params.loc[feat, target]
                    direction = directions[target].get(feat, "~")
                    direction_text.loc[feat, target] = f"{val:.2f}\n{direction}"
        except Exception:
            for feat in top_10_params.index:
                for target in top_10_params.columns:
                    val = top_10_params.loc[feat, target]
                    direction_text.loc[feat, target] = f"{val:.2f}\n~"
        
        # Визуализация heatmap
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.heatmap(top_10_params, annot=direction_text, fmt='', cmap='YlOrRd', cbar_kws={'label': 'Важность'}, ax=ax, linewidths=0.5, annot_kws={'size':9,'va':'center'})
        ax.set_title('Топ 10 параметров по влиянию на механические свойства (сталь)\n(со стрелками направления влияния)', fontsize=14, fontweight='bold')
        ax.set_ylabel('Параметры', fontsize=12)
        ax.set_xlabel('Целевые показатели', fontsize=12)
        st.pyplot(fig)
    
    uploaded_test = st.file_uploader("Загрузите Excel с новыми плавками (только химия)", type=['xlsx', 'xls'], key="t1")
    
    if uploaded_test:
        df_test = engineer_features(pd.read_excel(uploaded_test))
        X_base = df_test[base_features].fillna(df_test[base_features].mean())
        
        # Массовый прогноз для всех строк
        for t in targets:
            df_test[f'ПРОГНОЗ {t}'] = np.round(st.session_state.base_models[t].predict(X_base), 2)
            
        st.success(f"Прогноз выполнен для {len(df_test)} плавок!")
        st.dataframe(df_test[['Номер плавки'] + [f'ПРОГНОЗ {t}' for t in targets] + base_features] if 'Номер плавки' in df_test.columns else df_test)
        
        # Скачивание в Excel
        buf_ex = io.BytesIO()
        df_test.to_excel(buf_ex, index=False)
        st.download_button("📥 Скачать полную таблицу (Excel)", buf_ex.getvalue(), "predictions_steel.xlsx")
        
        st.subheader("Оптимизатор параметров плавки")
        if 'Номер плавки' in df_test.columns:
            selected_heat_num = st.selectbox("Выберите номер плавки для оптимизации:", df_test['Номер плавки'].values)
            selected_row = df_test[df_test['Номер плавки'] == selected_heat_num].index[0]
        else:
            selected_row = st.selectbox("Выберите индекс плавки для оптимизации:", df_test.index)
        
        # Анализ параметров с негативным влиянием на KV / K1C при увеличении
        row_base = X_base.iloc[[selected_row]].copy()
        baseline_preds = {target: st.session_state.base_models[target].predict(row_base)[0] for target in targets}
        negative_effects = []
        all_effects = []
        for feature in base_features:
            if feature not in row_base.columns:
                continue
            current_val_feat = float(row_base.iloc[0][feature])
            delta_val = abs(current_val_feat) * 0.1 if abs(current_val_feat) > 1e-6 else 0.1
            up_row = row_base.copy()
            up_row[feature] = current_val_feat + delta_val
            preds_up = {target: st.session_state.base_models[target].predict(up_row)[0] for target in targets}
            diffs = {target: preds_up[target] - baseline_preds[target] for target in targets}
            kv_dir = '↓' if diffs['KV Среднее'] < 0 or diffs['KV мин'] < 0 else '↑'
            k1c_dir = '↓' if diffs['К1С Среднее'] < 0 or diffs['К1С мин'] < 0 else '↑'
            row = {
                'Параметр': feature,
                'Текущее значение': current_val_feat,
                'Δ KV Среднее': diffs['KV Среднее'],
                'Δ KV мин': diffs['KV мин'],
                'Δ К1С Среднее': diffs['К1С Среднее'],
                'Δ К1С мин': diffs['К1С мин'],
                'KV влияние': kv_dir,
                'K1C влияние': k1c_dir
            }
            all_effects.append(row)
            if kv_dir == '↓' or k1c_dir == '↓':
                negative_effects.append(row)
        
        if negative_effects:
            st.subheader('Параметры, оказывающие негативное влияние при увеличении')
            neg_df = pd.DataFrame(negative_effects)
            neg_df[['Δ KV Среднее','Δ KV мин','Δ К1С Среднее','Δ К1С мин']] = neg_df[['Δ KV Среднее','Δ KV мин','Δ К1С Среднее','Δ К1С мин']].round(2)
            st.dataframe(neg_df)
        else:
            st.info('На этой плавке при увеличении базовых параметров негативных влияний на KV/K1C не выявлено.')
        
        param_to_opt = st.selectbox("Какой параметр хотите изменить?", base_features)
        current_val = float(df_test.loc[selected_row, param_to_opt])
        new_val = st.number_input("Введите новое значение:", value=current_val, format="%.4f")
        
        if st.button("Оценить влияние"):
            X_row = X_base.iloc[[selected_row]].copy()
            X_row[param_to_opt] = new_val
            st.write(f"**Результат для плавки {selected_row} при {param_to_opt} = {new_val}:**")
            cols = st.columns(4)
            for i, t in enumerate(targets):
                old_pred = df_test.loc[selected_row, f'ПРОГНОЗ {t}']
                new_pred = st.session_state.base_models[t].predict(X_row)[0]
                cols[i].metric(t, f"{new_pred:.2f}", f"{new_pred - old_pred:.2f}")

# =========================================================================
# ФУНКЦИЯ 2: СОВЕТНИК ПО ПРОКАТУ
# =========================================================================
elif mode == "2. Советник по прокату (Рекомендация режимов)":
    st.header("Функция 2: Подбор идеального режима термообработки")
    
    # Показываем важность параметров для расширенной модели (со сталью + прокат)
    with st.expander("📊 Анализ влияния параметров на механические свойства (сталь + прокат)", expanded=False):
        st.info("📈 **↑** = повышение параметра улучшает показатель | **↓** = повышение параметра ухудшает показатель | **~** = слабое влияние")
        
        # Собираем важности параметров из всех расширенных моделей
        importance_dict_ext = {}
        for target in targets:
            model = st.session_state.extended_models[target]
            importance_dict_ext[target] = model.feature_importances_
        
        # Создаём DataFrame с важностями
        importance_df_ext = pd.DataFrame(importance_dict_ext, index=extended_features)
        
        # Отдельно показываем влияние параметров ПРОКАТА (термообработки)
        st.subheader("Параметры проката/термообработки:")
        
        # Берём только параметры проката
        termo_importance = importance_df_ext.loc[termo_features].copy()
        
        # Рассчитываем направление влияния для проката
        try:
            # Создаём df для расчёта средних значений
            df_dummy_ext = pd.DataFrame({col: [0] for col in extended_features})
            for col in extended_features:
                df_dummy_ext[col] = np.random.uniform(0, 1, 100)
            
            directions_ext = calculate_feature_direction(st.session_state.extended_models, extended_features, df_dummy_ext)
            
            # Добавляем направления в DataFrame для аннотаций
            direction_text_ext = pd.DataFrame(index=termo_importance.index, columns=termo_importance.columns, dtype=object)
            for feat in termo_importance.index:
                for target in termo_importance.columns:
                    val = termo_importance.loc[feat, target]
                    direction = directions_ext[target].get(feat, "~")
                    direction_text_ext.loc[feat, target] = f"{val:.2f}\n{direction}"
        except Exception:
            direction_text_ext = pd.DataFrame(index=termo_importance.index, columns=termo_importance.columns, dtype=str)
            for feat in termo_importance.index:
                for target in termo_importance.columns:
                    val = termo_importance.loc[feat, target]
                    direction_text_ext.loc[feat, target] = f"{val:.2f}\n~"
        
        # Визуализация heatmap для параметров проката
        fig, ax = plt.subplots(figsize=(12, 3))
        sns.heatmap(termo_importance, annot=direction_text_ext, fmt='', cmap='RdYlGn', cbar_kws={'label': 'Важность'}, ax=ax, linewidths=0.5, annot_kws={'size':9,'va':'center'})
        ax.set_title('Параметры проката/термообработки и их влияние на механические свойства', fontsize=14, fontweight='bold')
        ax.set_ylabel('Параметры проката', fontsize=12)
        ax.set_xlabel('Целевые показатели', fontsize=12)
        st.pyplot(fig)
        
        # Полная карта: топ 10 всех параметров
        st.subheader("Топ 10 всех параметров (сталь + прокат) по важности:")
        
        importance_df_ext['Mean'] = importance_df_ext.mean(axis=1)
        top_10_params_ext = importance_df_ext.nlargest(10, 'Mean').drop('Mean', axis=1)
        
        try:
            direction_text_ext_full = pd.DataFrame(index=top_10_params_ext.index, columns=top_10_params_ext.columns, dtype=object)
            for feat in top_10_params_ext.index:
                for target in top_10_params_ext.columns:
                    val = top_10_params_ext.loc[feat, target]
                    direction = directions_ext[target].get(feat, "~")
                    direction_text_ext_full.loc[feat, target] = f"{val:.2f}\n{direction}"
        except Exception:
            direction_text_ext_full = pd.DataFrame(index=top_10_params_ext.index, columns=top_10_params_ext.columns, dtype=str)
            for feat in top_10_params_ext.index:
                for target in top_10_params_ext.columns:
                    val = top_10_params_ext.loc[feat, target]
                    direction_text_ext_full.loc[feat, target] = f"{val:.2f}\n~"
        
        fig, ax = plt.subplots(figsize=(12, 6))
        sns.heatmap(top_10_params_ext, annot=direction_text_ext_full, fmt='', cmap='RdYlGn', cbar_kws={'label': 'Важность'}, ax=ax, linewidths=0.5, annot_kws={'size':9,'va':'center'})
        ax.set_title('Топ 10 всех параметров по влиянию на механические свойства', fontsize=14, fontweight='bold')
        ax.set_ylabel('Параметры', fontsize=12)
        ax.set_xlabel('Целевые показатели', fontsize=12)
        st.pyplot(fig)
    
    uploaded_test = st.file_uploader("Загрузите Excel с плавками (химия)", type=['xlsx', 'xls'], key="t2")
    
    if uploaded_test:
        df_test = engineer_features(pd.read_excel(uploaded_test))
        if 'Номер плавки' in df_test.columns:
            selected_heat_num = st.selectbox("Выберите номер плавки для подбора термообработки:", df_test['Номер плавки'].unique())
            selected_row = df_test[df_test['Номер плавки'] == selected_heat_num].index[0]
        else:
            selected_row = st.selectbox("Выберите плавку для подбора термообработки:", df_test.index)
        
        if st.button("Запустить симулятор подбора режимов", type="primary"):
            with st.spinner("Создаю сотни виртуальных сценариев и ищу оптимум..."):
                # Сетка режиммов для симулятора
                heat_grid = np.linspace(200, 500, 7) # мин
                quench_grid = np.linspace(150, 300, 7) # сек
                temp_grid = np.linspace(500, 600, 7) # град
                
                grid = list(product(heat_grid, quench_grid, temp_grid))
                sim_df = pd.DataFrame(grid, columns=termo_features)
                
                # Добавляем химию выбранной плавки
                for col in base_features:
                    sim_df[col] = df_test.loc[selected_row, col] if col in df_test.columns else df_test[col].mean()
                if 'Mn/S' not in sim_df.columns: sim_df['Mn/S'] = df_test.loc[selected_row, 'Mn/S']
                
                X_sim = sim_df[extended_features].fillna(0)
                
                # Предсказываем
                for t in targets:
                    sim_df[f'P_{t}'] = st.session_state.extended_models[t].predict(X_sim)
                
                # Применяем ограничения: минимальные пороги для каждого параметра
                constraints = {
                    'P_KV мин': 10,        # не ниже 10
                    'P_KV Среднее': 10,    # не ниже 10
                    'P_К1С мин': 75,       # не ниже 75
                    'P_К1С Среднее': 82    # не ниже 82
                }
                
                # Фильтруем режимы, которые удовлетворяют всем ограничениям
                valid_regimes = sim_df.copy()
                for param, min_val in constraints.items():
                    valid_regimes = valid_regimes[valid_regimes[param] >= min_val]
                
                if len(valid_regimes) == 0:
                    st.warning("⚠️ Не найдено режимов, удовлетворяющих всем требованиям. Показываем лучший из доступных.")
                    valid_regimes = sim_df.copy()
                
                # Создаём комплексный Score: максимизируем все 4 параметра одновременно
                # Нормализуем каждый параметр и берём их взвешенную сумму
                valid_regimes['Score'] = (
                    (valid_regimes['P_KV мин'] / valid_regimes['P_KV мин'].max()) * 0.25 +
                    (valid_regimes['P_KV Среднее'] / valid_regimes['P_KV Среднее'].max()) * 0.25 +
                    (valid_regimes['P_К1С мин'] / valid_regimes['P_К1С мин'].max()) * 0.25 +
                    (valid_regimes['P_К1С Среднее'] / valid_regimes['P_К1С Среднее'].max()) * 0.25
                )
                
                best_idx = valid_regimes['Score'].idxmax()
                best_regime = sim_df.loc[best_idx]
                
            st.success("✅ Оптимальный режим найден!")
            st.write("### Рекомендуемые режимы проката/термообработки:")
            c1, c2, c3 = st.columns(3)
            c1.metric("Нагрев", f"{int(best_regime[termo_features[0]])} мин")
            c2.metric("Закалка", f"{int(best_regime[termo_features[1]])} с")
            c3.metric("Отпуск", f"{int(best_regime[termo_features[2]])} °C")
            
            st.write("### Ожидаемые механические свойства при этом режиме:")
            c4, c5, c6, c7 = st.columns(4)
            
            # Проверяем соответствие каждого параметра требованиям
            kv_min_val = best_regime['P_KV мин']
            kv_avg_val = best_regime['P_KV Среднее']
            k1c_min_val = best_regime['P_К1С мин']
            k1c_avg_val = best_regime['P_К1С Среднее']
            
            # Показываем с индикатором выполнения требований
            c4.metric("KV мин", f"{kv_min_val:.2f}", delta=f"(мин: 10)" if kv_min_val >= 10 else "(❌ ниже 10)")
            c5.metric("KV Среднее", f"{kv_avg_val:.2f}", delta=f"(мин: 10)" if kv_avg_val >= 10 else "(❌ ниже 10)")
            c6.metric("К1С мин", f"{k1c_min_val:.2f}", delta=f"(мин: 75)" if k1c_min_val >= 75 else "(❌ ниже 75)")
            c7.metric("К1С Среднее", f"{k1c_avg_val:.2f}", delta=f"(мин: 82)" if k1c_avg_val >= 82 else "(❌ ниже 82)")
            
            # Проверка выполнения всех требований
            all_met = (kv_min_val >= 10 and kv_avg_val >= 10 and k1c_min_val >= 75 and k1c_avg_val >= 82)
            if all_met:
                st.success("✅ Все требования выполнены!")
            else:
                st.warning("⚠️ Некоторые требования не выполнены. Это лучший доступный режим из симуляции.")
            
            # Экспорт в Word
            doc = Document()
            doc.add_heading('Рекомендация по термообработке', 0)
            doc.add_paragraph(f"Анализ для плавки индекс: {selected_row}")
            doc.add_heading('Режимы:', level=1)
            doc.add_paragraph(f"Нагрев: {int(best_regime[termo_features[0]])} мин\nЗакалка: {int(best_regime[termo_features[1]])} с\nОтпуск: {int(best_regime[termo_features[2]])} °C")
            doc.add_heading('Ожидаемые свойства:', level=1)
            for t in targets: doc.add_paragraph(f"{t}: {best_regime[f'P_{t}']:.2f}")
            
            buf_w = io.BytesIO()
            doc.save(buf_w)
            st.download_button("📄 Скачать Отчет (Word)", buf_w.getvalue(), "thermo_recommendation.docx")

# =========================================================================
# ФУНКЦИЯ 3: КОМБИНИРОВАННАЯ (ПОЛНЫЙ АНАЛИЗ)
# =========================================================================
elif mode == "3. Комбинированная (Полный анализ)":
    st.header("Функция 3: Оценка всей цепочки и рекомендации по улучшению")
    uploaded_test = st.file_uploader("Загрузите Excel (химия + фактический прокат)", type=['xlsx', 'xls'], key="t3")
    
    if uploaded_test:
        df_test = engineer_features(pd.read_excel(uploaded_test))
        X_ext = df_test[extended_features].fillna(df_test[extended_features].mean())
        
        for t in targets:
            df_test[f'ПРОГНОЗ {t}'] = np.round(st.session_state.extended_models[t].predict(X_ext), 2)
            
        st.success("Массовый прогноз выполнен!")
        st.dataframe(df_test[['Номер плавки'] + [f'ПРОГНОЗ {t}' for t in targets] + extended_features] if 'Номер плавки' in df_test.columns else df_test)
        
        st.subheader("Поиск 'узких мест' (Анализатор)")
        if 'Номер плавки' in df_test.columns:
            selected_heat_num = st.selectbox("Выберите номер плавки для глубокого анализа:", df_test['Номер плавки'].unique())
            selected_row = df_test[df_test['Номер плавки'] == selected_heat_num].index[0]
        else:
            selected_row = st.selectbox("Выберите плавку для глубокого анализа:", df_test.index)
        
        if st.button("Проанализировать параметры"):
            base_row = X_ext.iloc[[selected_row]].copy()
            current_preds = {t: st.session_state.extended_models[t].predict(base_row)[0] for t in targets}
            thresholds = {'KV мин': 10, 'KV Среднее': 10, 'К1С мин': 75, 'К1С Среднее': 82}
            
            st.success("✅ Анализ параметров завершен.")
            current_cols = st.columns(4)
            current_cols[0].metric("KV мин", f"{current_preds['KV мин']:.2f}", f"целевое ≥ {thresholds['KV мин']}")
            current_cols[1].metric("KV Среднее", f"{current_preds['KV Среднее']:.2f}", f"целевое ≥ {thresholds['KV Среднее']}")
            current_cols[2].metric("К1С мин", f"{current_preds['К1С мин']:.2f}", f"целевое ≥ {thresholds['К1С мин']}")
            current_cols[3].metric("К1С Среднее", f"{current_preds['К1С Среднее']:.2f}", f"целевое ≥ {thresholds['К1С Среднее']}")
            
            improvements = []
            for col in extended_features:
                if col == 'Mn/S':
                    continue
                current_val = float(base_row.iloc[0][col])
                candidates = []
                for action, factor in [('уменьшить', 0.9), ('увеличить', 1.1)]:
                    new_val = current_val * factor
                    test_row = base_row.copy()
                    test_row[col] = new_val
                    new_preds = {t: st.session_state.extended_models[t].predict(test_row)[0] for t in targets}
                    deltas = {t: new_preds[t] - current_preds[t] for t in targets}
                    weight = {t: 2.0 if current_preds[t] < thresholds[t] else 1.0 for t in targets}
                    score = sum(weight[t] * max(deltas[t], 0) for t in targets) - 0.5 * sum(abs(min(deltas[t], 0)) for t in targets)
                    candidates.append((action, new_val, deltas, score))
                best_action, best_val, best_deltas, best_score = max(candidates, key=lambda x: x[3])
                if best_score > 0:
                    improvements.append({
                        'Параметр': col,
                        'Рекомендация': f"{best_action} на 10% до {best_val:.2f}",
                        'Δ KV мин': best_deltas['KV мин'],
                        'Δ KV Среднее': best_deltas['KV Среднее'],
                        'Δ К1С мин': best_deltas['К1С мин'],
                        'Δ К1С Среднее': best_deltas['К1С Среднее'],
                        'Суммарный балл': best_score
                    })
            
            steel_improvements = sorted(
                [item for item in improvements if item['Параметр'] in base_features and item['Параметр'] != 'Mn/S'],
                key=lambda x: x['Суммарный балл'],
                reverse=True
            )[:5]
            rolling_improvements = sorted(
                [item for item in improvements if item['Параметр'] in termo_features],
                key=lambda x: x['Суммарный балл'],
                reverse=True
            )[:5]
            
            if steel_improvements:
                st.subheader('Топ-5 параметров производства стали для улучшения')
                steel_df = pd.DataFrame(steel_improvements)
                steel_df[['Δ KV мин', 'Δ KV Среднее', 'Δ К1С мин', 'Δ К1С Среднее', 'Суммарный балл']] = steel_df[['Δ KV мин', 'Δ KV Среднее', 'Δ К1С мин', 'Δ К1С Среднее', 'Суммарный балл']].round(2)
                st.dataframe(steel_df)
            else:
                st.warning('Не найдено параметров производства стали, улучшение которых прогнозируется как полезное для KV/K1C.')
            
            if rolling_improvements:
                st.subheader('Рекомендуемые параметры проката/термообработки')
                rolling_df = pd.DataFrame(rolling_improvements)
                rolling_df[['Δ KV мин', 'Δ KV Среднее', 'Δ К1С мин', 'Δ К1С Среднее', 'Суммарный балл']] = rolling_df[['Δ KV мин', 'Δ KV Среднее', 'Δ К1С мин', 'Δ К1С Среднее', 'Суммарный балл']].round(2)
                st.dataframe(rolling_df)
            else:
                st.warning('Не найдено параметров проката, улучшение которых прогнозируется как полезное для KV/K1C.')
            
            if improvements:
                best_imp = max(improvements, key=lambda x: x['Суммарный балл'])
                st.info(f"**Главное узкое место:** {best_imp['Параметр']} — {best_imp['Рекомендация']}, ожидаемый синтетический прирост: {best_imp['Суммарный балл']:.2f}")
            else:
                st.info('Параметры не показали положительного синтетического эффекта на KV/K1C при изменении на 10%.')

# =========================================================================
# ФУНКЦИЯ 4: РУЧНОЙ СИМУЛЯТОР
# =========================================================================
elif mode == "4. Ручной Симулятор (Песочница)":
    st.header("Функция 4: Интерактивный симулятор (Песочница)")
    st.write("Изменяйте параметры ползунками и смотрите, как реагирует металл.")
    
    col_s1, col_s2 = st.columns(2)
    inputs = {}
    
    with col_s1:
        st.subheader("Сталеплавильные параметры")
        inputs['О'] = st.slider("Кислород (О), ppm", 1.0, 10.0, 2.5)
        inputs['H'] = st.slider("Водород (H), ppm", 1.0, 3.0, 1.5)
        inputs['Mn'] = st.slider("Марганец (Mn), %", 0.5, 1.0, 0.75)
        inputs['S'] = st.slider("Сера (S), %", 0.001, 0.02, 0.005)
        inputs['Cr'] = st.slider("Хром (Cr), %", 0.1, 0.5, 0.25)
        inputs['N'] = st.slider("Азот (N), %", 0.001, 0.02, 0.008)
        inputs['Al'] = st.slider("Алюминий (Al), %", 0.01, 0.05, 0.02)
        inputs['As'] = st.slider("Мышьяк (As), %", 0.001, 0.02, 0.005)
        inputs['Ti'] = st.slider("Титан (Ti), %", 0.001, 0.02, 0.005)
        inputs['Перегрев МНЛЗ(конечный)'] = st.slider("Перегрев конечный, °C", 10.0, 50.0, 25.0)
        inputs['Перегрев МНЛЗ (начальный)'] = st.slider("Перегрев начальный, °C", 10.0, 60.0, 30.0)
        
    with col_s2:
        st.subheader("Параметры термообработки")
        inputs['Продолжительность нагрева, мин'] = st.slider("Время нагрева, мин", 100.0, 600.0, 300.0)
        inputs['Факт время закалки ВЗМ, с'] = st.slider("Время закалки, с", 100.0, 400.0, 200.0)
        inputs['темп-ра отпуска, °С'] = st.slider("Температура отпуска, °C", 400.0, 650.0, 500.0)
        
        st.divider()
        st.subheader("Мгновенный прогноз:")
        
        inputs['Mn/S'] = inputs['Mn'] / inputs['S']
        sim_df = pd.DataFrame([inputs])
        
        c1, c2 = st.columns(2)
        c3, c4 = st.columns(2)
        
        c1.metric("KV мин", f"{st.session_state.extended_models['KV мин'].predict(sim_df[extended_features])[0]:.2f}")
        c2.metric("KV Среднее", f"{st.session_state.extended_models['KV Среднее'].predict(sim_df[extended_features])[0]:.2f}")
        c3.metric("К1С мин", f"{st.session_state.extended_models['К1С мин'].predict(sim_df[extended_features])[0]:.2f}")
        c4.metric("К1С Среднее", f"{st.session_state.extended_models['К1С Среднее'].predict(sim_df[extended_features])[0]:.2f}")